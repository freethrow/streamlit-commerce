# process.py

# opn the Italy - Serbia excel file 

import pandas as pd

import docx

from docx.shared import Pt, Inches

import plotly.express as px

from words import words

def growth(num):

    return round(num-100,1)

# translate the names of commodity groups
def translate(name):
    if name in words:
        return words[name]
    return name


# function to process the excel file into a dataframe
def process(data):
    
    """
        Returns all of the data in a dataframe.
    """
    
    # drop NA values and replace * and - with zeros
    data.dropna(inplace=True)
    data.replace('*',0, inplace=True)
    data.replace('-',0, inplace=True)
    data['Voce'] = data['Naziv'].str.strip()
    data['Voce'] = data['Voce'].apply(lambda x:translate(x))

    data['Var. Export'] = pd.to_numeric(data['Indeks - izvoz'])
    data['Var. Import'] = pd.to_numeric(data['Indeks - uvoz'])
    
    data['Var. Export'] = data['Var. Export'].apply(lambda x:growth(x))
    data['Var. Import'] = data['Var. Import'].apply(lambda x:growth(x))

    data['Esportazioni'] = pd.to_numeric(data['Izvoz'], downcast='integer')
    data['Importazioni'] = pd.to_numeric(data['Uvoz'], downcast='integer')
    data['Interscambio'] = data['Esportazioni']+data['Importazioni']
    data['Surplus'] = data['Esportazioni']-data['Importazioni']   

    data = data[['Voce','Esportazioni','Importazioni','Var. Export','Var. Import','Interscambio','Surplus']]
    
    return data


# read the file
original = pd.read_excel('data/settoriItaliaAprile.xls',skiprows=2,usecols="B:F")


# process the data
df_data = process(original)

# reindexed
df_data_r = df_data

print(df_data)

# get the totals
totals = df_data.loc[df_data['Voce'] == 'TOTALE']
w_totals = totals.set_index('Voce', inplace=False)




# open an existing document
doc = docx.Document()

doc.add_picture('logo.png', width=Inches(2.5))
doc.add_heading('Interscambio Serbia - Italia', 0)

# add a table to the end and create a reference variable
# extra row is so we can add the header row
t = doc.add_table(df_data_r.shape[0]+1, df_data_r.shape[1], style='Light Grid Accent 5')

# add the header rows.
for j in range(df_data_r.shape[-1]):
    t.cell(0,j).text = df_data_r.columns[j]

# add the rest of the data frame
for i in range(df_data_r.shape[0]):
    for j in range(df_data_r.shape[-1]):
        t.cell(i+1,j).text = str(df_data_r.values[i,j])
        
        
for row in t.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run_obj = paragraph.runs
        run = run_obj[0]
        font = run.font
        font.size = Pt(8)
        
doc.add_page_break()


# save the doc
doc.save('test1.docx')
